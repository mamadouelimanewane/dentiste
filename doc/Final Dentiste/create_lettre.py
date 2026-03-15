#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Lettre_Presentation_DentoPrestige_Final.docx
Modeled exactly after "lettre presentation.docx" for LexPremium
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Lettre_Presentation_DentoPrestige_Final.docx")

# Colors
GOLD_RGB   = RGBColor(0xD4, 0xAF, 0x37)
DARK_RGB   = RGBColor(0x0F, 0x17, 0x2A)
SLATE_RGB  = RGBColor(0x33, 0x41, 0x55)
WHITE_RGB  = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_heading(doc, text, level=1, color=GOLD_RGB, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size or (14 if level == 1 else 12))
    run.font.bold = True
    run.font.color.rgb = color
    return p

def shade_paragraph(p, hex_color="0f172a"):
    """Apply paragraph background shading."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def add_hr(doc, color_hex="D4AF37"):
    """Add a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── HEADER BLOCK ────────────────────────────────────────────────────────
    # Brand name
    p_brand = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2)
    r_logo = p_brand.add_run("\u25c6  DENTOPRESTIGE")
    r_logo.font.name = "Calibri"
    r_logo.font.size = Pt(22)
    r_logo.font.bold = True
    r_logo.font.color.rgb = GOLD_RGB

    p_tag = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6)
    r_tag = p_tag.add_run("L'\u00c9cosyst\u00e8me d\u2019IA & Haute Gestion pour Cabinets Dentaires d\u2019Exception")
    r_tag.font.name = "Calibri"
    r_tag.font.size = Pt(11)
    r_tag.font.italic = True
    r_tag.font.color.rgb = SLATE_RGB

    add_hr(doc)

    # Date & destinataire
    p_date = add_para(doc, space_before=10, space_after=4)
    r_date = p_date.add_run("Dakar, le 13 Mars 2026")
    r_date.font.name = "Calibri"
    r_date.font.size = Pt(11)
    r_date.font.color.rgb = SLATE_RGB

    add_para(doc, space_before=0, space_after=2)
    p_dest = add_para(doc, space_before=0, space_after=2)
    r_dest = p_dest.add_run("\u00c0 l\u2019attention de [TITRE] [NOM PR\u00c9NOM]")
    r_dest.font.name = "Calibri"; r_dest.font.size = Pt(11); r_dest.font.bold = True

    p_cabinet = add_para(doc, space_before=0, space_after=2)
    r_cab = p_cabinet.add_run("[NOM DU CABINET DENTAIRE]")
    set_font(r_cab)

    p_addr = add_para(doc, space_before=0, space_after=8)
    r_addr = p_addr.add_run("[ADRESSE]")
    set_font(r_addr)

    # Object
    p_obj = add_para(doc, space_before=6, space_after=8)
    r_obj_lbl = p_obj.add_run("Objet : ")
    r_obj_lbl.font.name = "Calibri"; r_obj_lbl.font.size = Pt(11); r_obj_lbl.font.bold = True; r_obj_lbl.font.color.rgb = GOLD_RGB
    r_obj_txt = p_obj.add_run("Pr\u00e9sentation de DentoPrestige, Solution innovante de gestion pour cabinets dentaires")
    set_font(r_obj_txt)

    add_hr(doc)

    # ── SALUTATION ─────────────────────────────────────────────────────────
    p_sal = add_para(doc, space_before=10, space_after=8)
    r_sal = p_sal.add_run("Docteur,")
    set_font(r_sal, bold=True, size=11)

    # ── CORPS — INTRODUCTION ────────────────────────────────────────────────
    p_intro = add_para(doc, space_before=0, space_after=8)
    r_intro = p_intro.add_run(
        "Dans un contexte o\u00f9 40\u0025 de votre temps est absorb\u00e9 par l\u2019administratif, "
        "j\u2019ai l\u2019honneur de vous pr\u00e9senter DentoPrestige, une solution r\u00e9volutionnaire "
        "sp\u00e9cialement con\u00e7ue pour les cabinets dentaires de l\u2019espace OHADA. "
        "Notre plateforme ne se contente pas de digitaliser vos processus \u2014 elle les transforme "
        "radicalement."
    )
    set_font(r_intro)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── CORPS — POURQUOI ───────────────────────────────────────────────────
    add_heading(doc, "Pourquoi DentoPrestige est unique pour votre cabinet ?", level=1)

    sections = [
        (
            "Conformit\u00e9 OHADA Native",
            "Contrairement aux solutions occidentales inadapt\u00e9es, DentoPrestige int\u00e8gre "
            "nativement le SYSCOHADA, la gestion FSE (Feuille de Soins \u00c9lectronique) et les "
            "normes sanitaires UEMOA. Chaque module respecte scrupuleusement les r\u00e9glementations "
            "m\u00e9dicales et comptables de notre espace.",
        ),
        (
            "Intelligence Artificielle Clinique Avanc\u00e9e",
            "Notre Scanner Radiologique analyse en 60 secondes vos clich\u00e9s et d\u00e9tecte "
            "automatiquement les pathologies. L\u2019IA Clinique \u00e9value vos plans de traitement "
            "sur la base de l\u2019historique patient. Le G\u00e9n\u00e9rateur de Protocoles cr\u00e9e "
            "des plans de soins en 3 clics avec un gain de temps de 75\u0025.",
        ),
        (
            "\u00c9cosyst\u00e8me Complet sur Mesure",
            "44+ modules v4.0 couvrent l\u2019int\u00e9gralit\u00e9 de vos besoins : "
            "gestion de dossiers patients intelligente, Holo-Smile Studio AR, Neural Vision Lab, "
            "Blockchain Medical Ledger, Messagerie Interne, Concierge Bio-Logistique, "
            "GED avec OCR haute r\u00e9solution, comptabilit\u00e9 analytique en temps r\u00e9el.",
        ),
        (
            "Adaptation aux R\u00e9alit\u00e9s Locales",
            "Int\u00e9gration native de Wave et Orange Money, communication unifi\u00e9e "
            "WhatsApp/SMS/Email, fonctionnement optimal m\u00eame avec connexion instable, "
            "interface en fran\u00e7ais m\u00e9dical OHADA.",
        ),
    ]

    for i, (title, body) in enumerate(sections):
        p_num = doc.add_paragraph()
        p_num.paragraph_format.space_before = Pt(6)
        p_num.paragraph_format.space_after  = Pt(2)
        p_num.paragraph_format.left_indent  = Cm(0.5)
        r_n = p_num.add_run(f"{i+1}.  ")
        r_n.font.bold = True; r_n.font.color.rgb = GOLD_RGB; r_n.font.size = Pt(11); r_n.font.name = "Calibri"
        r_t = p_num.add_run(title)
        r_t.font.bold = True; r_t.font.size = Pt(11); r_t.font.name = "Calibri"; r_t.font.color.rgb = DARK_RGB

        p_b = add_para(doc, space_before=0, space_after=6)
        p_b.paragraph_format.left_indent = Cm(1.0)
        r_b = p_b.add_run(body)
        set_font(r_b)
        p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── CORPS — RÉSULTATS ─────────────────────────────────────────────────
    add_heading(doc, "Les r\u00e9sultats concrets pour votre cabinet :", level=1)

    resultats = [
        "R\u00e9cup\u00e9ration de 40\u0025 de votre temps administratif pour vous concentrer sur le soin strat\u00e9gique",
        "Am\u00e9lioration significative de votre taux de recouvrement gr\u00e2ce aux relances automatiques",
        "Z\u00e9ro rendez-vous manqu\u00e9 avec les alertes intelligentes de rappels patients",
        "Transparence totale avec vos patients via le portail extranet s\u00e9curis\u00e9",
        "Conformit\u00e9 m\u00e9dico-l\u00e9gale absolue avec tra\u00e7abilit\u00e9 compl\u00e8te des consentements",
    ]

    for res in resultats:
        p_r = add_para(doc, space_before=2, space_after=2)
        p_r.paragraph_format.left_indent = Cm(0.5)
        r_bullet = p_r.add_run("\u25c6  ")
        r_bullet.font.color.rgb = GOLD_RGB; r_bullet.font.bold = True; r_bullet.font.name = "Calibri"; r_bullet.font.size = Pt(11)
        r_res = p_r.add_run(res)
        set_font(r_res)

    add_para(doc, space_before=4, space_after=4)

    # ── CORPS — DÉMO & CTA ────────────────────────────────────────────────
    p_lead = add_para(doc, space_before=8, space_after=6)
    r_lead = p_lead.add_run(
        "Soyez le cabinet leader de Dakar et de la sous-r\u00e9gion qui aura franchi le pas "
        "et constater une transformation profonde de votre efficacit\u00e9 op\u00e9rationnelle. "
        "La version de d\u00e9monstration compl\u00e8te est accessible d\u00e8s maintenant sur "
    )
    set_font(r_lead)
    p_lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_url = p_lead.add_run("https://dentoprestige.vercel.app/")
    r_url.font.name = "Calibri"; r_url.font.size = Pt(11); r_url.font.color.rgb = RGBColor(0x22, 0x86, 0xD4)
    r_creds = p_lead.add_run(" (identifiants : admin@dentoprestige.sn / demo2026).")
    set_font(r_creds)

    p_rdv = add_para(doc, space_before=6, space_after=6)
    r_rdv = p_rdv.add_run(
        "Je sollicite l\u2019honneur d\u2019un rendez-vous \u00e0 votre convenance pour une "
        "d\u00e9monstration personnalis\u00e9e de 45 minutes au cours de laquelle je vous "
        "pr\u00e9senterai concr\u00e8tement comment DentoPrestige peut transformer la gestion "
        "de votre cabinet et vous faire gagner des dizaines d\u2019heures chaque mois."
    )
    set_font(r_rdv)
    p_rdv.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── FERMETURE ─────────────────────────────────────────────────────────
    p_close = add_para(doc, space_before=10, space_after=6)
    r_close = p_close.add_run(
        "Dans l\u2019attente de votre retour, je vous prie d\u2019agr\u00e9er, Docteur, "
        "l\u2019expression de ma haute consid\u00e9ration."
    )
    set_font(r_close)
    p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_para(doc, space_before=16, space_after=2)

    # Signature block
    for line, bold, color, size in [
        ("Mamadou Dia",                    True,  DARK_RGB,  12),
        ("Directeur G\u00e9n\u00e9ral",    False, SLATE_RGB, 11),
        ("ProcessIng\u00e9nierie",          True,  GOLD_RGB,  11),
        ("T\u00e9l : +221 777 529 288",     False, SLATE_RGB, 10),
        ("Email : mamadou.dia@processingenierie.com", False, SLATE_RGB, 10),
    ]:
        p_s = add_para(doc, space_before=1, space_after=1)
        r_s = p_s.add_run(line)
        r_s.font.name = "Calibri"; r_s.font.size = Pt(size); r_s.font.bold = bold; r_s.font.color.rgb = color

    add_hr(doc)

    p_footer = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
    r_f = p_footer.add_run("DentoPrestige Africa  \u2022  Soignez mieux, g\u00e9rez moins  \u2022  Le Standard de l\u2019Excellence pour le Praticien de Demain")
    r_f.font.name = "Calibri"; r_f.font.size = Pt(8); r_f.font.italic = True; r_f.font.color.rgb = GOLD_RGB

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")

if __name__ == "__main__":
    main()
